"""各格式解析器 —— 全部产出统一的 ParsedDoc（PLAN §12.2）。

设计原则：**解析失败不能拖垮流水线**。单个文档加密、损坏、格式异常时，
返回已解析出的部分 + warnings，而不是抛异常让整批索引中断。
只有完全无法读取才抛 ParseError。
"""

from __future__ import annotations

import re
import zipfile
from html.parser import HTMLParser
from pathlib import Path, PurePosixPath

from app.parsing.base import (
    Block,
    BlockKind,
    Locator,
    ParsedDoc,
    ParseError,
    UnsupportedFormat,
)

# 单文档解析上限：超大文档解析会吃满内存且价值递减
MAX_CHARS = 5_000_000

# DOCX is a ZIP container. A small compressed upload can otherwise make
# python-docx allocate hundreds of MiB before the shared text limit is reached.
DOCX_MAX_ARCHIVE_BYTES = 16 * 1024 * 1024
DOCX_MAX_ENTRIES = 2_000
DOCX_MAX_ENTRY_BYTES = 32 * 1024 * 1024
DOCX_MAX_TOTAL_BYTES = 128 * 1024 * 1024
DOCX_MAX_COMPRESSION_RATIO = 300
_DOCX_RATIO_MIN_BYTES = 1024 * 1024


def parse_pdf(path: Path) -> ParsedDoc:
    import pymupdf as fitz

    blocks: list[Block] = []
    warnings: list[str] = []

    try:
        doc = fitz.open(path)
    except Exception as e:
        raise ParseError(f"无法打开 PDF：{e}") from e

    try:
        if doc.needs_pass:
            raise ParseError("PDF 已加密")

        total = 0
        for pno in range(doc.page_count):
            try:
                page = doc.load_page(pno)
                text = page.get_text("text")
            except Exception:
                warnings.append(f"第 {pno + 1} 页解析失败，已跳过")
                continue

            for para in _split_paragraphs(text):
                blocks.append(
                    Block(
                        kind=BlockKind.PARAGRAPH,
                        text=para,
                        locator=Locator(page=pno + 1),
                    )
                )
                total += len(para)

            if total > MAX_CHARS:
                warnings.append(f"文档过长，只索引了前 {pno + 1} 页")
                break

        page_count = doc.page_count
        title = (doc.metadata or {}).get("title") or None
    finally:
        doc.close()

    # 扫描件（纯图片 PDF）没有文本层 → 走当前平台系统 OCR（可在设置关闭）
    if not blocks:
        from app.parsing import ocr

        if ocr.runtime_enabled() and ocr.is_available():
            pages, ocr_warnings = ocr.ocr_pdf(path)
            warnings.extend(ocr_warnings)
            for page_no, text in pages:
                for para in _split_paragraphs(text):
                    blocks.append(
                        Block(
                            kind=BlockKind.PARAGRAPH,
                            text=para,
                            locator=Locator(page=page_no),
                        )
                    )
            if blocks:
                warnings.append(f"扫描件：已通过 OCR 提取 {len(pages)} 页文本")
            else:
                warnings.append("未提取到文本（OCR 也没有识别出内容）")
        else:
            warnings.append("未提取到文本，可能是扫描件（OCR 未启用或不可用）")

    return ParsedDoc(blocks=blocks, title=title, page_count=page_count, warnings=warnings)


def _validate_docx_archive(path: Path) -> None:
    """Reject encrypted, ambiguous and explosively compressed DOCX archives."""
    try:
        archive_size = path.stat().st_size
    except OSError as exc:
        raise ParseError(f"无法读取 DOCX：{exc}") from exc
    if archive_size > DOCX_MAX_ARCHIVE_BYTES:
        raise ParseError("DOCX 压缩包超过 16 MiB 解析上限")

    try:
        with zipfile.ZipFile(path) as archive:
            entries = archive.infolist()
            if len(entries) > DOCX_MAX_ENTRIES:
                raise ParseError(f"DOCX ZIP 条目过多（上限 {DOCX_MAX_ENTRIES}）")
            total = 0
            names: set[str] = set()
            has_document = False
            for entry in entries:
                name = entry.filename.replace("\\", "/")
                parts = PurePosixPath(name).parts
                if (not name or name.startswith("/") or ".." in parts
                        or name in names):
                    raise ParseError("DOCX ZIP 包含非法或重复条目")
                names.add(name)
                if entry.flag_bits & 0x1:
                    raise ParseError("DOCX ZIP 包含加密条目")
                if entry.file_size > DOCX_MAX_ENTRY_BYTES:
                    raise ParseError("DOCX ZIP 单个条目解压后过大")
                total += entry.file_size
                if total > DOCX_MAX_TOTAL_BYTES:
                    raise ParseError("DOCX ZIP 解压总量过大")
                if entry.file_size >= _DOCX_RATIO_MIN_BYTES:
                    ratio = entry.file_size / max(1, entry.compress_size)
                    if ratio > DOCX_MAX_COMPRESSION_RATIO:
                        raise ParseError("DOCX ZIP 压缩比异常，已拒绝解析")
                has_document = has_document or name == "word/document.xml"
            if not has_document:
                raise ParseError("DOCX 缺少 word/document.xml")
    except ParseError:
        raise
    except (OSError, zipfile.BadZipFile, zipfile.LargeZipFile) as exc:
        raise ParseError(f"无法打开 DOCX ZIP：{exc}") from exc


def parse_docx(path: Path) -> ParsedDoc:
    import docx  # python-docx

    _validate_docx_archive(path)
    try:
        doc = docx.Document(str(path))
    except Exception as e:
        raise ParseError(f"无法打开 DOCX：{e}") from e

    blocks: list[Block] = []
    heading_path: list[str] = []
    warnings: list[str] = []
    total_chars = 0

    def append_bounded(block: Block) -> bool:
        nonlocal total_chars
        remaining = MAX_CHARS - total_chars
        if remaining <= 0:
            if not warnings:
                warnings.append(f"文档过长，只索引前 {MAX_CHARS} 个字符")
            return False
        if len(block.text) > remaining:
            block.text = block.text[:remaining]
            warnings.append(f"文档过长，只索引前 {MAX_CHARS} 个字符")
        if block.text:
            blocks.append(block)
            total_chars += len(block.text)
        return total_chars < MAX_CHARS

    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        style = (para.style.name or "").lower() if para.style else ""
        level = _heading_level(style)

        if level:
            heading_path = heading_path[: level - 1] + [text]
            keep_going = append_bounded(
                Block(
                    kind=BlockKind.HEADING,
                    text=text,
                    level=level,
                    locator=Locator(para_index=idx, heading_path=list(heading_path)),
                )
            )
        else:
            kind = BlockKind.LIST_ITEM if "list" in style else BlockKind.PARAGRAPH
            keep_going = append_bounded(
                Block(
                    kind=kind,
                    text=text,
                    locator=Locator(para_index=idx, heading_path=list(heading_path)),
                )
            )
        if not keep_going:
            break

    # 表格单独成块，保留行结构（分片器会整表保留，见 §12.2c）
    for table in doc.tables:
        if total_chars >= MAX_CHARS:
            break
        rows = []
        for row in table.rows:
            remaining = MAX_CHARS - total_chars - sum(len(value) for value in rows)
            if remaining <= 0:
                break
            cells = []
            for cell in row.cells:
                value = cell.text.strip()
                if not value:
                    cells.append("")
                    continue
                value = value[:remaining]
                cells.append(value)
                remaining -= len(value) + 3
                if remaining <= 0:
                    break
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            if not append_bounded(
                Block(
                    kind=BlockKind.TABLE,
                    text="\n".join(rows),
                    locator=Locator(heading_path=list(heading_path)),
                )
            ):
                break

    title = blocks[0].text if blocks and blocks[0].kind == BlockKind.HEADING else None
    return ParsedDoc(blocks=blocks, title=title, warnings=warnings)


_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")
_MD_FENCE = re.compile(r"^```")


def parse_markdown(path: Path) -> ParsedDoc:
    text = _read_text(path)
    blocks: list[Block] = []
    heading_path: list[str] = []
    buf: list[str] = []
    in_code = False
    code_buf: list[str] = []

    def flush_para():
        if buf:
            content = "\n".join(buf).strip()
            if content:
                blocks.append(
                    Block(
                        kind=BlockKind.PARAGRAPH,
                        text=content,
                        locator=Locator(heading_path=list(heading_path)),
                    )
                )
            buf.clear()

    for line in text.splitlines():
        if _MD_FENCE.match(line):
            if in_code:
                blocks.append(
                    Block(
                        kind=BlockKind.CODE,
                        text="\n".join(code_buf),
                        locator=Locator(heading_path=list(heading_path)),
                    )
                )
                code_buf.clear()
            else:
                flush_para()
            in_code = not in_code
            continue

        if in_code:
            code_buf.append(line)
            continue

        m = _MD_HEADING.match(line)
        if m:
            flush_para()
            level = len(m.group(1))
            title = m.group(2).strip()
            heading_path = heading_path[: level - 1] + [title]
            blocks.append(
                Block(
                    kind=BlockKind.HEADING,
                    text=title,
                    level=level,
                    locator=Locator(heading_path=list(heading_path)),
                )
            )
        elif not line.strip():
            flush_para()
        else:
            buf.append(line)

    if in_code and code_buf:  # 代码块未闭合
        blocks.append(
            Block(kind=BlockKind.CODE, text="\n".join(code_buf),
                  locator=Locator(heading_path=list(heading_path)))
        )
    flush_para()

    title = blocks[0].text if blocks and blocks[0].kind == BlockKind.HEADING else None
    return ParsedDoc(blocks=blocks, title=title)


def parse_text(path: Path) -> ParsedDoc:
    text = _read_text(path)
    blocks = [
        Block(kind=BlockKind.PARAGRAPH, text=p, locator=Locator(para_index=i))
        for i, p in enumerate(_split_paragraphs(text))
    ]
    return ParsedDoc(blocks=blocks)


class _TextExtractor(HTMLParser):
    """从 HTML 抽纯文本 —— 标准库实现，零依赖，PyInstaller 友好。

    面向两类：微信/QQ 导出的聊天记录 HTML，以及保存的网页。
    丢弃 script/style，块级标签边界断行，<br> 换行，标题另存供文档标题。
    """

    _BLOCK = {"p", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6",
              "section", "article", "blockquote", "td", "th"}
    _SKIP = {"script", "style", "noscript", "svg"}

    def __init__(self):
        super().__init__()
        self.parts: list[str] = []
        self.title: str | None = None
        self._skip_depth = 0
        self._in_title = False

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip_depth += 1
        elif tag == "title":
            self._in_title = True
        elif tag == "br" or tag in self._BLOCK:
            self.parts.append("\n")

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip_depth:
            self._skip_depth -= 1
        elif tag == "title":
            self._in_title = False
        elif tag in self._BLOCK:
            self.parts.append("\n")

    def handle_data(self, data):
        if self._skip_depth:
            return
        if self._in_title:
            self.title = (self.title or "") + data.strip()
        elif data.strip():
            self.parts.append(data)


def parse_html(path: Path) -> ParsedDoc:
    import html as _html

    extractor = _TextExtractor()
    try:
        extractor.feed(_read_text(path))
    except Exception:
        pass  # 残缺 HTML 也尽量抽已解析的部分
    text = _html.unescape("".join(extractor.parts))
    blocks = [
        Block(kind=BlockKind.PARAGRAPH, text=p, locator=Locator(para_index=i))
        for i, p in enumerate(_split_paragraphs(text))
    ]
    title = (extractor.title or "").strip() or None
    warnings = [] if blocks else ["未从 HTML 中提取到文本"]
    return ParsedDoc(blocks=blocks, title=title, warnings=warnings)


PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_text,
    ".rtf": parse_text,
    ".html": parse_html,
    ".htm": parse_html,
    ".csv": parse_text,
}


def parse(path: Path | str) -> ParsedDoc:
    p = Path(path)
    parser = PARSERS.get(p.suffix.lower())
    if parser is None:
        raise UnsupportedFormat(f"不支持的格式：{p.suffix}")
    return parser(p)


# ---------------------------------------------------------------- 内部工具

def _read_text(path: Path) -> str:
    """读文本文件。中文文档常见 GBK 编码，UTF-8 失败时回退。"""
    raw = path.read_bytes()[: MAX_CHARS * 4]
    for enc in ("utf-8", "gb18030", "utf-16", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _split_paragraphs(text: str) -> list[str]:
    """按空行切段，并合并被硬换行拆散的行。

    PDF 提取的文本常在固定宽度处换行，直接按行切会把一句话切碎。
    """
    parts = re.split(r"\n\s*\n", text)
    out = []
    for part in parts:
        merged = re.sub(r"\n(?![\s•\-\d])", "", part).strip()
        if merged:
            out.append(merged)
    return out


def _heading_level(style_name: str) -> int:
    """从 DOCX 样式名判断标题级别。中英文样式名都要认。"""
    m = re.search(r"heading\s*(\d)", style_name)
    if m:
        return int(m.group(1))
    m = re.search(r"标题\s*(\d)", style_name)
    if m:
        return int(m.group(1))
    if style_name in ("title", "标题"):
        return 1
    return 0
